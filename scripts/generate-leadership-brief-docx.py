# -*- coding: utf-8 -*-
"""生成智慧业财系统领导汇报 Word 文档"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "智慧业财系统-领导汇报材料.docx"


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text, size=12, bold=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, name="黑体", size=16 if level == 1 else 14, bold=True)
    return h


def add_bullets(doc, items, size=12):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=size)
        p.paragraph_format.space_after = Pt(3)


def add_numbered(doc, items, size=12):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run, size=size)
        p.paragraph_format.space_after = Pt(3)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True, size=11)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=11)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)

    # 封面
    add_para(doc, "萍乡再生资源智慧业财系统", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(doc, "平台作用与核心模块说明", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_para(doc, "（呈报公司领导 · 经营财务部 / 基地管理部 / 业务经营部 已上线使用）", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, "文档性质：汇报说明材料，可根据汇报场合自行增删", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=48)
    doc.add_page_break()

    # 一、摘要
    add_heading(doc, "一、摘要：平台解决什么问题", 1)
    add_para(
        doc,
        "智慧业财系统（算账经营平台）将废钢再生资源业务的「采购—生产—销售—核算」贯通到同一套数据中，"
        "让经营财务部、基地管理部（生产）、业务经营部（销售）在同一事实基础上看成本、看利润、做决策。"
        "系统已上线并投入日常使用，本材料重点说明平台价值及成本分析、加工单录入、利润分析三大模块，"
        "其中利润核算以材料成本（LIFO 溯源）为核心。",
    )
    add_bullets(
        doc,
        [
            "采购视角：花了多少钱、买了什么料、库存值多少（成本分析）。",
            "生产视角：每天加工多少成品、消耗哪些毛料（加工单录入）。",
            "销售视角：每一张发货结算单赚不赚钱、赚在哪里亏在哪里（利润分析）。",
            "管理视角：按日/周/月汇总，支持导出 Excel，明细可追溯到单、到批、到料型。",
        ],
    )

    # 二、平台定位
    add_heading(doc, "二、平台定位与业务主线", 1)
    add_heading(doc, "2.1 平台定位", 2)
    add_para(
        doc,
        "平台面向废钢铁等再生资源经营场景，不是简单的报表展示，而是把业务单据与财务核算规则嵌入系统，"
        "实现「业务发生即留下核算依据、销售完成即可逐单算利」。",
    )
    add_table(
        doc,
        ["模块分组", "主要功能", "主要使用部门"],
        [
            ["算账经营", "成本分析、利润分析、生产加工录入、运维管理", "经营财务部、基地管理部、业务经营部"],
            ["首页数据", "废钢铁/报废车等交易明细汇总", "业务、管理"],
            ["能碳管理", "能耗、碳排放、固废等", "能碳管理相关部门"],
            ["处置采购", "竞价、一口价、各类废弃物采购", "采购业务"],
        ],
    )

    add_heading(doc, "2.2 业务数据主线（建议向领导展示的逻辑链）", 2)
    add_numbered(
        doc,
        [
            "采购入库：毛料进入基地，形成采购成本与库存（PurchaseWarehouse）。",
            "生产加工：按日录入成品产量与各毛料消耗，系统同步核减毛料、增加成品库存（ProcessingCostInput）。",
            "销售结算：客户发货结算，形成含税收入、结算量、净重等（DeliverySettlement）。",
            "材料成本反推：对每一张销售单，按后进先出（LIFO）原则，从加工记录中溯源该成品的毛料构成与成本（MaterialCostCache / 实时 LIFO）。",
            "利润核算：在材料成本基础上，叠加加工成本、运输费、税费、贴现、回款利息及即征即退、政府扶持等，得到单据级利润。",
        ],
    )
    add_para(doc, "一句话概括：先记清楚「进了什么料、产出了什么货、卖了多少钱」，再按规则逐单算清楚「这笔货的材料成本从哪来、最终赚没赚钱」。", bold=True)

    # 三、成本分析
    add_heading(doc, "三、成本分析模块（花钱视角）", 1)
    add_para(doc, "入口：算账经营 → 成本分析。口号：实时监控废钢采购成本，为经营决策提供数据支持。")
    add_heading(doc, "3.1 模块作用", 2)
    add_para(
        doc,
        "成本分析只回答一个问题：采购环节花了多少钱、花在了哪里、结构是否合理。"
        "数据全部来自采购入库表，不掺杂销售与利润逻辑，便于与财务采购账、基地收货台账对照。",
    )
    add_heading(doc, "3.2 领导可关注的指标", 2)
    add_bullets(
        doc,
        [
            "当日 / 最近一周 / 本月采购成本（万元）及对应吨数。",
            "按业务类型拆分：基地收货（SH）、基地买货（TH）、协同业务。",
            "最近 30 日成本日趋势、最近一周结构分解。",
            "当月各废钢料型成本占比及平均采购单价，并与上月对比。",
            "毛料库存价值分析：当前库存吨数、加权采购单价、库存总金额。",
            "支持按日期导出毛料采购汇总、基地收货与成品统计台账。",
        ],
    )
    add_heading(doc, "3.3 对管理的价值", 2)
    add_bullets(
        doc,
        [
            "发现异常高价采购日、料型结构突变。",
            "为销售定价、利润测算提供「原料成本底座」参考。",
            "与基地生产、销售节奏联动，解释「为什么某阶段利润波动」。",
        ],
    )

    # 四、加工单录入
    add_heading(doc, "四、加工单录入模块（生产视角）", 1)
    add_para(doc, "入口：算账经营 → 生产加工录入（Web）；基地现场亦可通过小程序写入同一数据表。")
    add_heading(doc, "4.1 模块作用", 2)
    add_para(
        doc,
        "加工单是利润核算的「生产事实来源」。系统按生产日期、成品名称、成品库区（如 A1/B1/C1/D2），"
        "记录当日成品产量（吨）以及各毛料的消耗重量与单价。一次提交同时完成：",
    )
    add_bullets(
        doc,
        [
            "毛料库存核减（MaterialStorage）：按录入的材料构成扣减对应库区毛料。",
            "成品库存增加（ProductStock）：成品进入对应库区。",
            "为后续 LIFO 材料成本计算提供「哪批货、哪天产、用了哪些料、每吨料多少钱」的明细依据。",
        ],
    )
    add_heading(doc, "4.2 录入要点（与利润准确性直接相关）", 2)
    add_bullets(
        doc,
        [
            "成品名称、成品库区必须与销售结算单上的品种/库别能够对应，否则销售单无法匹配到生产批次。",
            "生产日期不得晚于该成品的销售发货日期。",
            "各毛料列（如 MSLKM、MGJKM 等）须填写用量与单价；缺项会导致该批材料成本偏低甚至为 0。",
            "毛料总消耗重量应不小于成品产量（系统有校验）；成本明显超收入时会提示确认。",
        ],
    )
    add_para(
        doc,
        "向领导说明：加工单录入质量决定利润分析中「材料成本」是否可信；生产条线的及时、准确录入，是财务逐单算利的前提。",
        bold=True,
    )

    # 五、利润分析
    add_heading(doc, "五、利润分析模块（赚钱视角）", 1)
    add_para(doc, "入口：算账经营 → 利润分析。口号：实时监控废钢业务利润，为经营决策提供数据支持。")
    add_heading(doc, "5.1 核算粒度", 2)
    add_para(
        doc,
        "利润分析以每一张销售结算单（发货单）为最小核算单元，而不是只做月度粗汇总。"
        "领导既可以看到今日/本周/本月利润总览，也可以下钻到任意一单，查看收入、材料成本、加工成本、"
        "其它成本、其它收入及吨钢毛利。",
    )
    add_heading(doc, "5.2 页面展示的利润公式", 2)
    add_para(doc, "系统界面采用的利润公式为：", bold=True)
    add_para(
        doc,
        "利润 = 销售收入÷1.13（折不含税）"
        " − 材料成本 − 加工成本"
        " − 运输费 − 税费 + 即征即退 + 政府扶持资金"
        " − 贴现费用 − 回款周期资金利息",
    )
    add_table(
        doc,
        ["项目", "含义", "说明"],
        [
            ["销售收入（含税）", "DeliverySettlement 结算金额", "按单取数"],
            ["材料成本", "LIFO 反推的废钢（毛料）成本", "核心项，见第六节"],
            ["加工成本", "加工费单价 × 净重（或结算量）", "按成品配置"],
            ["其它成本", "运输费 + 税费 + 贴现 + 回款利息", "按客户与参数表计算"],
            ["其它收入", "即征即退 + 政府扶持", "按客户与参数表计算"],
            ["吨钢毛利", "利润 ÷ 出厂净重", "元/吨，便于横向比较"],
        ],
    )
    add_heading(doc, "5.3 领导可见的分析内容", 2)
    add_bullets(
        doc,
        [
            "汇总卡片：今日/本周/本月利润、今日材料成本与加工成本。",
            "日利润趋势（近 30 天）、最近一周利润分解柱状图。",
            "销售明细利润分析表：支持按月份筛选、分页、导出 Excel。",
            "成品对比：按「客户—成品」维度，对比当月与上月销量及平均含税单价。",
            "明细行悬停可查看：材料 LIFO 构成、税费/运输/贴现等子项的计算过程与参数快照。",
        ],
    )

    # 六、LIFO 核心
    add_heading(doc, "六、材料成本核算原则（核心）：LIFO 与毛料溯源", 1)
    add_para(
        doc,
        "向领导汇报时建议强调：本系统利润核算的「科学性」主要体现在材料成本——"
        "不是按一个固定吨价或财务手工分摊，而是按每一张销售单，结合真实生产加工记录，"
        "采用后进先出（LIFO）原则，反推该单成品所消耗的毛料构成与金额。",
        bold=True,
    )

    add_heading(doc, "6.1 为什么材料成本是核心", 2)
    add_bullets(
        doc,
        [
            "再生资源业务中，销售收入相对透明，但原料采购批次多、价格波动大、混料加工常见。",
            "若材料成本口径不准，利润、税费基数、即征即退及政府扶持等联动项都会失真。",
            "因此系统把「材料成本怎么算」作为利润模型的地基，其余成本/收入项在材料成本确定后按规则叠加。",
        ],
    )

    add_heading(doc, "6.2 LIFO（后进先出）原则说明", 2)
    add_para(
        doc,
        "当某一成品对外销售时，系统假设：优先消耗「最近生产」的批次来匹配该笔销量（与物理上的堆垛发运习惯一致）。"
        "具体步骤如下：",
    )
    add_numbered(
        doc,
        [
            "识别销售单对应的成品：根据结算单上的成品名称（product_type）与库区/warehouse，匹配加工表 ProcessingCostInput 中同成品、同库区的生产记录。",
            "筛选可用批次：只取生产日期 ≤ 发货日期的批次，且当日产量 dailyProcess_qty > 0。",
            "按生产日期从新到旧排序：先扣最近一天的生产 batch，再扣更早 batch，直至满足该单的「材料核算量」。",
            "确定核算量：优先用结算量，并结合净重、磅差（transitloss）或路损系数；避免整车净重误套在拆分结算单上。",
            "计算各批材料成本：每一生产批次中，各毛料成本 = 毛料用量 × 毛料单价；批次单位材料成本 = 各毛料成本之和 ÷ 成品产量；按 LIFO 摊到本销售单上的吨数，累加得到该单材料成本总额。",
            "输出溯源结果：除总金额外，还保存 material_composition（各毛料吨数与金额）及 production_records（使用了哪些生产日、哪几条加工记录），供财务与管理复核。",
        ],
    )

    add_heading(doc, "6.3 示意（文字版流程）", 2)
    add_para(doc, "销售单 FHxxxx → 成品「PG钢筋压块」、发货日 6月2日、核算量 34.51 吨", bold=True)
    add_para(doc, "↓ 匹配加工记录（同成品、同库区，生产日期 ≤ 6月2日）")
    add_para(doc, "↓ 按日期倒序：先取 6月1日批次 20 吨 → 再取 5月31日批次 14.51 吨 …")
    add_para(doc, "↓ 每批次按录入的毛料用量×单价，算出该批每吨材料成本")
    add_para(doc, "↓ 汇总 → 本单材料成本 = Σ(各批摊销吨数 × 该批单位材料成本)")
    add_para(doc, "↓ 写入 MaterialCostCache，利润分析读取并展示；悬停可见毛料构成明细")

    add_heading(doc, "6.4 入库单税率溯源（与税费、扶持联动）", 2)
    add_para(
        doc,
        "材料成本不仅影响「减项」，还影响税费、即征即退、政府扶持等项目的「基数」。"
        "系统对 LIFO 溯源得到的各毛料，进一步从采购入库记录 PurchaseWarehouse 中，"
        "查找该毛料在生产日期之前最近一笔入库的 tax_rate（入库单税率），"
        "按材料用量加权平均，得到该销售单的「入库单加权税率」。"
        "若无法溯源，则回退至参数表中的默认入库单税率。",
    )

    add_heading(doc, "6.5 缓存与运维保障", 2)
    add_bullets(
        doc,
        [
            "优先读取 MaterialCostCache（按发货单号缓存），加快利润分析加载。",
            "缓存未命中或为 0 时，系统自动实时 LIFO 重算。",
            "运维控制台提供「刷新材料成本缓存」等工具，供财务在批量导入或规则调整后统一重算。",
            "加工单、采购单补录或修正后，应触发缓存刷新，以保证利润数据与业务事实一致。",
        ],
    )

    # 七、其它成本收入
    add_heading(doc, "七、其它成本与收入项（参数化、按客户区分）", 1)
    add_para(doc, "在材料成本、加工成本确定后，系统按 ProfitParamConfig 参数表逐单计算：")
    add_table(
        doc,
        ["类别", "主要项目", "适用说明"],
        [
            ["其它成本", "运输费", "按客户配置运价、路损系数 × 净重"],
            ["其它成本", "税费", "基于统一基数×税率；基数含收入、材料、加工、运输等"],
            ["其它成本", "贴现费用", "主要为萍钢客户配置"],
            ["其它成本", "回款周期资金利息", "按客户配置回款天数与年利率"],
            ["其它收入", "即征即退", "主要为新钢客户"],
            ["其它收入", "政府扶持资金", "萍钢/吉钢/新钢等按政策比例配置"],
        ],
    )
    add_para(
        doc,
        "税费、即征即退、政府扶持共用的「基数」口径为："
        "收入不含税×13% − 材料成本×入库单税率 − 加工成本×9% − 运输费×3%。"
        "参数可在运维页维护，改参后利润分析按发货日期取生效值。",
    )

    # 八、部门协同
    add_heading(doc, "八、已上线部门的协同分工", 1)
    add_table(
        doc,
        ["部门", "在系统中的主要工作", "对利润准确性的贡献"],
        [
            ["经营财务部", "采购入库维护、销售结算、参数维护、利润/成本报表分析、导出对账", "保证收入与采购成本数据完整；审核利润异常"],
            ["基地管理部（生产）", "每日加工单录入（产量、毛料消耗）", "提供 LIFO 材料成本溯源的生产依据"],
            ["业务经营部（销售）", "关注销售明细、客户—成品对比、吨钢毛利", "反馈市场与结算异常，驱动数据修正"],
        ],
    )

    # 九、汇报建议
    add_heading(doc, "九、向领导汇报时的建议话术（可参考）", 1)
    add_numbered(
        doc,
        [
            "这套系统已经把采购、生产、销售串成一条链，三个部门在用的就是同一套数。",
            "成本分析看「钱怎么花出去」，利润分析看「每一单货赚没赚钱」。",
            "利润不是拍脑袋分摊，核心是材料成本：按 LIFO 从加工记录里把毛料构成倒推出来，能追到单、追到批。",
            "加工单录得越及时越准，利润就越可信；财务可以在明细里点开看每一单的计算过程。",
            "后续可在此基础上做客户结构优化、料型采购决策、定价与返利政策评估。",
        ],
    )

    # 十、附录
    add_heading(doc, "附录：主要数据表对照（供完善文档时参考）", 1)
    add_table(
        doc,
        ["数据表", "业务含义"],
        [
            ["PurchaseWarehouse", "采购入库 / 毛料采购成本"],
            ["ProcessingCostInput", "生产加工录入 / 产量与毛料消耗"],
            ["DeliverySettlement", "销售发货结算 / 收入与销量"],
            ["MaterialCostCache", "按销售单缓存的 LIFO 材料成本"],
            ["MaterialStorage", "毛料库存"],
            ["ProductStock", "成品库存"],
            ["ProfitParamConfig", "利润核算参数（运价、税率、贴现等）"],
            ["ProcessingCostConfig", "各成品加工费单价"],
        ],
    )

    add_para(doc, "—— 文档结束 ——", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(
        doc,
        "说明：本材料依据系统实际上线功能整理，具体界面与公式以平台最新版本为准。"
        "您可在 Word 中直接修改措辞、插入截图或补充本公司制度要求。",
        size=10,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"已生成: {OUTPUT}")


if __name__ == "__main__":
    build()
